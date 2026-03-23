import fitz  # PyMuPDF - fixed import
import docx
import re
from typing import Dict, List, Any
from pathlib import Path
import streamlit as st

class EnhancedResumeParser:
    """Enhanced resume parser with comprehensive skill extraction and better parsing"""
    
    def __init__(self):
        # Will get skill categories from config when needed
        self.skill_categories = {}
        self._load_skill_categories()
        
        # Common resume sections for better parsing
        self.section_patterns = {
            'contact': [r'contact', r'personal', r'information'],
            'summary': [r'summary', r'profile', r'objective', r'about'],
            'experience': [r'experience', r'employment', r'work', r'career', r'professional'],
            'education': [r'education', r'academic', r'qualification', r'degree'],
            'skills': [r'skills', r'technical', r'competenc', r'expertise'],
            'projects': [r'projects', r'portfolio', r'achievements'],
            'certifications': [r'certification', r'license', r'credential']
        }
    
    def _load_skill_categories(self):
        """Load skill categories from config"""
        try:
            from config import SKILL_CATEGORIES
            self.skill_categories = SKILL_CATEGORIES
        except ImportError:
            # Fallback basic categories
            self.skill_categories = {
                'programming_languages': ['python', 'java', 'javascript', 'c++', 'sql'],
                'frameworks': ['react', 'django', 'flask', 'spring'],
                'databases': ['mysql', 'postgresql', 'mongodb'],
                'cloud': ['aws', 'azure', 'gcp']
            }

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Enhanced PDF text extraction with better formatting"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                # Extract text with formatting preservation
                page_text = page.get_text("text")
                
                # Clean up common PDF artifacts
                page_text = re.sub(r'\n\s*\n', '\n\n', page_text)  # Multiple newlines
                page_text = re.sub(r'\s+', ' ', page_text)  # Multiple spaces
                page_text = re.sub(r'([a-z])([A-Z])', r'\1 \2', page_text)  # CamelCase separation
                
                text += page_text + "\n"
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            st.error(f"PDF extraction error: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """Enhanced DOCX text extraction"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text.strip() + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
            
        except Exception as e:
            st.error(f"DOCX extraction error: {e}")
            return ""

    def extract_text(self, file_path: str) -> str:
        """Extract text from various file formats with enhanced processing"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        text = ""
        
        if extension == '.pdf':
            text = self.extract_text_from_pdf(str(file_path))
        elif extension in ['.docx', '.doc']:
            text = self.extract_text_from_docx(str(file_path))
        elif extension == '.txt':
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except Exception as e:
                st.error(f"Text file error: {e}")
                return ""
        else:
            st.error(f"Unsupported file format: {extension}")
            return ""
        
        # Post-process extracted text
        if text:
            text = self._clean_extracted_text(text)
        
        return text

    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Multiple blank lines
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs
        
        # Fix common extraction issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Split camelCase
        text = re.sub(r'([.!?])([A-Z])', r'\1 \2', text)  # Space after punctuation
        
        # Normalize bullet points
        text = re.sub(r'[â€¢â–ªâ–«â—¦â€£âƒ]', 'â€¢', text)
        
        # Remove page numbers and headers/footers (common patterns)
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            # Skip likely page numbers, headers, footers
            if re.match(r'^\d+$', line):  # Just a number
                continue
            if re.match(r'^Page \d+', line, re.IGNORECASE):
                continue
            if len(line) < 3:  # Very short lines
                continue
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)

    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Enhanced contact information extraction"""
        contact_info = {
            'emails': [],
            'phones': [],
            'linkedin': [],
            'github': [],
            'location': []
        }
        
        # Email extraction - more comprehensive
        email_patterns = [
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            r'email[:\s]*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',
        ]
        
        for pattern in email_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            contact_info['emails'].extend(matches)
        
        # Phone number extraction - improved patterns
        phone_patterns = [
            r'(\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
            r'(\+?[0-9]{1,3}[-.\s]?)?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}',
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    phone = ''.join(match)
                else:
                    phone = match
                
                # Clean and validate phone number
                clean_phone = re.sub(r'[^\d+]', '', phone)
                if 10 <= len(clean_phone) <= 15:  # Reasonable phone number length
                    contact_info['phones'].append(phone)
        
        # LinkedIn URL extraction
        linkedin_patterns = [
            r'linkedin\.com/in/[A-Za-z0-9\-]+',
            r'linkedin\.com/profile/[A-Za-z0-9\-]+',
        ]
        
        for pattern in linkedin_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            contact_info['linkedin'].extend(matches)
        
        # GitHub URL extraction
        github_patterns = [
            r'github\.com/[A-Za-z0-9\-]+',
        ]
        
        for pattern in github_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            contact_info['github'].extend(matches)
        
        # Location extraction (basic)
        location_patterns = [
            r'([A-Z][a-z]+,\s*[A-Z]{2})',  # City, State
            r'([A-Z][a-z]+\s*[A-Z][a-z]+,\s*[A-Z]{2})',  # City Name, State
        ]
        
        for pattern in location_patterns:
            matches = re.findall(pattern, text)
            contact_info['location'].extend(matches)
        
        # Remove duplicates and clean up
        for key in contact_info:
            contact_info[key] = list(set(contact_info[key]))
        
        return contact_info

    def extract_enhanced_skills(self, text: str) -> Dict[str, List[str]]:
        """Enhanced skill extraction using comprehensive categories"""
        text_lower = text.lower()
        extracted_skills = {}
        
        for category, skills in self.skill_categories.items():
            found_skills = []
            
            for skill in skills:
                # Use word boundaries for more accurate matching
                skill_pattern = r'\b' + re.escape(skill.lower()) + r'\b'
                if re.search(skill_pattern, text_lower):
                    found_skills.append(skill)
            
            if found_skills:
                extracted_skills[category] = found_skills
        
        return extracted_skills

    def extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Enhanced work experience extraction"""
        experience = []
        lines = text.split('\n')
        
        # Look for experience section
        in_experience_section = False
        experience_keywords = ['experience', 'employment', 'work history', 'career', 'professional']
        current_job = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check if we're entering experience section
            if any(keyword in line.lower() for keyword in experience_keywords):
                in_experience_section = True
                continue
            
            # Stop if we hit another section
            section_headers = ['education', 'skills', 'projects', 'certifications']
            if any(header in line.lower() for header in section_headers) and len(line) < 50:
                in_experience_section = False
                continue
            
            if in_experience_section or not experience:  # Always try to extract some experience
                # Patterns for job titles and companies
                job_patterns = [
                    # "Job Title at Company Name"
                    r'^([^,\n]+?)\s+at\s+([^,\n]+?)(?:\s*[|\-,]\s*(.+))?$',
                    # "Job Title - Company Name"
                    r'^([^,\n]+?)\s*[-â€“â€”]\s*([^,\n]+?)(?:\s*[|\-,]\s*(.+))?$',
                    # "Job Title, Company Name"
                    r'^([^,\n]+?),\s*([^,\n]+?)(?:\s*[|\-,]\s*(.+))?$',
                ]
                
                for pattern in job_patterns:
                    match = re.match(pattern, line)
                    if match:
                        groups = match.groups()
                        # Determine which is job title vs company
                        part1, part2 = groups[0].strip(), groups[1].strip()
                        duration = groups[2].strip() if groups[2] else ""
                        
                        # Simple heuristic: job titles often contain common job words
                        job_words = ['engineer', 'developer', 'manager', 'analyst', 'director',
                                     'specialist', 'coordinator', 'associate', 'senior', 'junior',
                                     'lead', 'principal', 'architect', 'consultant', 'intern']
                        
                        if any(word in part1.lower() for word in job_words):
                            job_title, company = part1, part2
                        else:
                            company, job_title = part1, part2
                        
                        # Validate reasonable length
                        if 3 <= len(job_title) <= 100 and 3 <= len(company) <= 100:
                            experience.append({
                                'position': job_title,
                                'company': company,
                                'duration': duration,
                                'description': ''
                            })
                            current_job = len(experience) - 1
                        break
                
                # If we have a current job, add description lines
            elif current_job is not None and len(experience) > current_job:
                if line.startswith('â€¢') or line.startswith('-') or line.startswith('*'):
                    if experience[current_job]['description']:
                        experience[current_job]['description'] += ' ' + line
                    else:
                        experience[current_job]['description'] = line
        
        return experience[:8]  # Limit to reasonable number

    def extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Enhanced education extraction"""
        education = []
        
        # Common degree patterns
        degree_patterns = [
            r'(bachelor|master|phd|doctorate|bsc|msc|ba|ma|mba|bs|ms)(?:\s+of|\s+in|\s+degree)?\s+([^,\n.]+?)(?:\s*,\s*([^,\n.]+?))?(?:\s*[,\n.]|\s*$)',
            r'(associate|diploma|certificate)\s+(?:of|in)?\s*([^,\n.]+?)(?:\s*,\s*([^,\n.]+?))?',
        ]
        
        for pattern in degree_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                degree_type = match[0].strip().title()
                field = match[1].strip().title() if match[1] else ""
                institution = match[2].strip().title() if len(match) > 2 and match[2] else ""
                
                if field:  # Only add if we have a field of study
                    education.append({
                        'degree': degree_type,
                        'field': field,
                        'institution': institution
                    })
        
        # Remove duplicates
        seen = set()
        unique_education = []
        for edu in education:
            key = (edu['degree'], edu['field'])
            if key not in seen:
                seen.add(key)
                unique_education.append(edu)
        
        return unique_education

    def parse_resume(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Enhanced resume parsing with comprehensive information extraction"""
        try:
            # Extract raw text
            raw_text = self.extract_text(file_path)
            
            if not raw_text.strip():
                return {"success": False, "error": "Could not extract text from resume"}
            
            # Extract structured information
            contact_info = self.extract_contact_info(raw_text)
            skills = self.extract_enhanced_skills(raw_text)
            experience = self.extract_experience(raw_text)
            education = self.extract_education(raw_text)
            
            # Calculate text statistics
            word_count = len(raw_text.split())
            char_count = len(raw_text)
            line_count = len([line for line in raw_text.split('\n') if line.strip()])
            
            # Extract sections for better organization
            sections = self._identify_sections(raw_text)
            
            return {
                'success': True,  # Important for hybrid system
                'filename': filename,
                'cleaned_text': raw_text,  # For compatibility
                'raw_text': raw_text,
                'contact_info': contact_info,
                'skills': skills,
                'experience': experience,
                'education': education,
                'sections': sections,
                'statistics': {
                    'word_count': word_count,
                    'char_count': char_count,
                    'line_count': line_count,
                    'skills_found': len(skills),
                    'experience_entries': len(experience),
                    'education_entries': len(education)
                },
                'word_count': word_count  # Backward compatibility
            }
            
        except Exception as e:
            st.error(f"Resume parsing error for {filename}: {e}")
            return {"success": False, "error": f"Failed to parse resume: {str(e)}"}

    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identify and extract different sections of the resume"""
        sections = {}
        lines = text.split('\n')
        current_section = None
        section_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            found_section = None
            for section_name, patterns in self.section_patterns.items():
                if any(re.search(pattern, line, re.IGNORECASE) for pattern in patterns):
                    if len(line) < 50:  # Likely a header, not content
                        found_section = section_name
                        break
            
            if found_section:
                # Save previous section
                if current_section and section_content:
                    sections[current_section] = '\n'.join(section_content)
                
                # Start new section
                current_section = found_section
                section_content = []
            elif current_section:
                section_content.append(line)
        
        # Save last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)
        
        return sections

# Alias for compatibility with hybrid system
ResumeParser = EnhancedResumeParser
